import streamlit as st
import os
import io
import json
import re
import bcrypt
from datetime import datetime
import time
import base64 # Import base64 for image encoding

# --- Firebase Imports ---
import firebase_admin
from firebase_admin import credentials, auth, firestore
from firebase_admin import exceptions # Import exceptions module for FirebaseError

# --- Google Drive Imports ---
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# --- AI & Document Processing Imports ---
from openai import OpenAI
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn # For font color


# --- Streamlit Page Configuration (MUST BE THE FIRST ST COMMAND) ---
st.set_page_config(
    page_title="SSO Consultants AI Recruitment",
    page_icon="sso_logo.png", # Set favicon for the browser tab
    layout="wide" # Keeping wide layout, but centering content within it
)

# --- Custom CSS for Styling ---
st.markdown(
    """
    <style>
    /* Global base styling - pure white background, pure black text by default */
    body {
        background-color: #FFFFFF; /* Pure white background */
        color: #000000 !important; /* Pure black for general text readability - CRITICAL */
        font-family: 'Inter', sans-serif;
        height: 100vh; /* Ensure body takes full viewport height */
        margin: 0;
        padding: 0;
    }
    /* Streamlit's main app container */
    .stApp {
        background-color: #FFFFFF;
        color: #000000;
        min-height: 100vh; /* Ensure it takes full viewport height for centering */
        display: flex;
        flex-direction: column;
        justify-content: flex-start; /* Let content align to start, specific containers will center */
        align-items: stretch; /* Allow children to stretch, don't force overall app content to center here */
    }
    /* Headers */
    h1, h2, h3, h4, h5, h6 {
        color: #4CAF50; /* Green for headers from config.toml */
    }
    /* Buttons */
    .stButton>button {
        background-color: #4CAF50; /* Green background for buttons from config.toml */
        color: white; /* White text for buttons */
        border-radius: 5px;
        border: none;
        padding: 10px 20px;
        font-size: 16px;
        cursor: pointer;
        transition: background-color 0.3s ease;
        
        /* New/modified properties for consistent sizing and centering text */
        width: 100%; /* Make button fill its immediate container (e.g., column) */
        height: 50px; /* Set a fixed ideal height for all buttons */
        display: flex; /* Use flexbox to center text vertically and horizontally */
        align-items: center;
        justify-content: center;
        text-align: center; /* Fallback for text alignment */
        white-space: normal; /* Allow text to wrap if it's too long for the button */
        word-break: break-word; /* Break words if necessary */
    }
    .stButton>button:hover {
        background-color: #66BB6A; /* Lighter green on hover */
    }
    /* Text Inputs, Text Areas, Select Boxes */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea, .stSelectbox>div>div {
        border: 1px solid #4CAF50; /* Green border */
        border-radius: 5px;
        padding: 8px;
        color: #000000; /* Black text */
        background-color: #FFFFFF; /* White background */
    }
    /* Placeholder text */
    .stTextInput>div>div>input::placeholder, .stTextArea>div>div>textarea::placeholder {
        color: #888888; /* Grey placeholder */
    }
    /* Specific Streamlit components by data-testid */
    .stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
        font-size:1.2rem;
        color: #4CAF50; /* Green tab headers */
    }
    /* Centering content for specific pages */
    .main-content-centered {
        display: flex;
        flex-direction: column;
        align-items: center;
        text-align: center;
        padding-top: 50px;
    }
    .main-content-centered .stTextInput, .main-content-centered .stButton {
        max-width: 400px; /* Limit width of input/buttons in centered view */
        width: 100%;
    }

    /* New CSS for the fixed top-left logo */
    .fixed-top-left-logo {
        position: fixed;
        top: 15px; /* Distance from the top */
        left: 20px; /* Distance from the left */
        z-index: 1000; /* Ensure it stays on top of other elements */
        height: 60px; /* Adjust height as needed */
        width: auto;
    }

    /* Adjusted .login-container to remove min-height and border/shadow */
    .login-container {
        display: flex;
        flex-direction: column;
        align-items: center;
        width: 100%;
        max-width: 650px; /* MODIFIED: Increased width from 550px to 650px */
        margin: auto; /* Centers the content within its flex parent horizontally */
        padding: 20px; /* Reduced padding */
        background-color: #FFFFFF;
        /* No border, no box-shadow */
    }
    .login-container .stTextInput, .login-container .stButton {
        width: 100%; /* Ensure inputs and buttons fill the container width */
    }

    /* --- ALIGNMENT RULES --- */

    /* Target the main content area (excluding sidebar) */
    /* This targets the 'section' element with class 'main' that Streamlit uses for the central content */
    section.main {
        display: flex;
        flex-direction: column;
        align-items: center; /* Center content horizontally within the main section */
        flex-grow: 1; /* Allow it to grow and fill available vertical space */
        width: 100%; /* Ensure it takes full width */
        padding: 0 20px; /* Add some horizontal padding to prevent content from touching edges */
    }

    /* Streamlit's internal blocks within the main section */
    /* MODIFIED: Removed align-items: center; from here to allow more control */
    [data-testid="stVerticalBlock"] {
        display: flex;
        flex-direction: column;
        /* align-items: center;  REMOVED this line */
        width: 100%; /* Ensure it takes full width of its parent */
    }

    /* Adjust Streamlit's root element to allow flex centering */
    /* This centers the entire stApp component on the page */
    #root > div:first-child {
        display: flex;
        flex-direction: column;
        justify-content: center; /* Center vertically */
        align-items: center; /* Center horizontally */
        min-height: 100vh;
        width: 100%;
    }
    
    /* Ensure forms are centered if they are not inside a flex-centered container (like login-container) */
    form:not(.login-container form) { 
        width: 100%;
        max-width: 500px; /* Adjust as desired for general forms */
        margin: 0 auto; /* Center the form itself */
    }

    /* Sidebar content centering */
    [data-testid="stSidebarContent"] {
        display: flex;
        flex-direction: column;
        align-items: center; /* Centers content horizontally within the sidebar */
        padding-top: 20px; /* Adjust as needed */
    }
    /* For images specifically inside sidebar to ensure centering */
    [data-testid="stSidebarContent"] img {
        display: block; /* Important for margin: auto to work */
        margin-left: auto;
        margin-right: auto;
    }

    /* --- NEW CSS FOR LEFT-ALIGNED FILE UPLOADERS --- */
    .left-aligned-content {
        width: 100%; /* Ensure it takes full width of its parent */
        display: flex;
        flex-direction: column;
        align-items: flex-start; /* Aligns children to the left */
    }
    /* Ensure file uploader widgets themselves also align left within this container */
    .left-aligned-content [data-testid="stFileUploader"] {
        width: 100%; /* Take full width of the left-aligned container */
    }
    /* This targets the label of the file uploader */
    .left-aligned-content [data-testid="stFileUploader"] label {
        text-align: left !important; /* Force label text to left align */
        width: 100%; /* Ensure label spans full width */
    }
    /* This targets the inner vertical block within the file uploader */
    .left-aligned-content [data-testid="stFileUploader"] > div > [data-testid="stVerticalBlock"] {
        align-items: flex-start !important; /* Force inner vertical blocks also align left - ADDED !important */
    }
    /* This specifically targets the upload button text within the file uploader */
    .left-aligned-content [data-testid="stFileUploader"] button div span {
        text-align: left !important;
        width: 100%;
        justify-content: flex-start; /* Align button content to start */
    }
    /* Ensure the button itself also aligns left within the uploader */
    .left-aligned-content [data-testid="stFileUploader"] button {
        align-self: flex-start; /* Align the button element itself to the left */
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- Session State Initialization ---
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'user_email' not in st.session_state:
    st.session_state['user_email'] = None
if 'user_uid' not in st.session_state:
    st.session_state['user_uid'] = None
if 'is_admin' not in st.session_state:
    st.session_state['is_admin'] = False
if 'login_mode' not in st.session_state: # New: Control login flow
    st.session_state['login_mode'] = 'choose_role' # Start with role selection
if 'is_admin_attempt' not in st.session_state: # New: Flag for current login attempt type
    st.session_state['is_admin_attempt'] = False # Default to user login attempt
if 'username' not in st.session_state:
    st.session_state['username'] = None
if 'has_set_username' not in st.session_state:
    st.session_state['has_set_username'] = False
if 'needs_username_setup' not in st.session_state: # Flag to explicitly trigger setup page
    st.session_state['needs_username_setup'] = False
if 'login_success' not in st.session_state:
    st.session_state['login_success'] = False
if 'current_admin_page' not in st.session_state:
    st.session_state['current_admin_page'] = 'reports'


# --- Firebase Initialization ---
# Ensure only one app instance is initialized
if not firebase_admin._apps:
    try:
        # Load the Firebase service account key from Streamlit secrets
        firebase_service_account_info = st.secrets["SERVICE_ACCOUNT_KEY"]

        # Ensure private_key is correctly formatted with actual newlines
        if isinstance(firebase_service_account_info, dict) and "private_key" in firebase_service_account_info:
            firebase_service_account_info["private_key"] = firebase_service_account_info["private_key"].replace('\\n', '\n')

        cred = credentials.Certificate(firebase_service_account_info)
        firebase_admin.initialize_app(cred) # Use firebase_admin.initialize_app
        db = firestore.client()
        st.success("Firebase initialized successfully.")

    except KeyError:
        st.error("Firebase SERVICE_ACCOUNT_KEY not found in Streamlit secrets! "
                 "Please add your Firebase service account JSON content to your app's secrets "
                 "on Streamlit Community Cloud under the key 'SERVICE_ACCOUNT_KEY'.")
        st.stop()
    except Exception as e:
        st.error(f"An unexpected error occurred during Firebase initialization: {e}")
        st.info("Please ensure your 'SERVICE_ACCOUNT_KEY' is valid and correctly formatted.")
        st.stop()
else:
    db = firestore.client()

# --- Google Drive Configuration ---
drive_service = None
try:
    # Load the Google Drive key from Streamlit secrets
    google_drive_key_info = st.secrets["GOOGLE_DRIVE_KEY"]

    # Ensure private_key is correctly formatted with actual newlines
    if isinstance(google_drive_key_info, dict) and "private_key" in google_drive_key_info:
        google_drive_key_info["private_key"] = google_drive_key_info["private_key"].replace('\\n', '\n')

    # Define the necessary scopes for Google Drive access
    SCOPES = ['https://www.googleapis.com/auth/drive'] # Scope for full Drive access

    # Create credentials from the service account info
    drive_credentials = service_account.Credentials.from_service_account_info(google_drive_key_info, scopes=SCOPES)

    # Build the Google Drive API service client
    drive_service = build('drive', 'v3', credentials=drive_credentials)
    st.success("Google Drive API initialized successfully.")

except KeyError:
    st.error("Google Drive GOOGLE_DRIVE_KEY not found in Streamlit secrets! "
             "Please add your Google Drive service account JSON content to your app's secrets "
             "on Streamlit Community Cloud under the key 'GOOGLE_DRIVE_KEY'.")
    st.stop()
except Exception as e:
    st.error(f"An unexpected error occurred during Google Drive initialization: {e}")
    st.info("Please ensure your 'GOOGLE_DRIVE_KEY' is valid and Google Drive API is enabled.")
    st.stop()

# --- OpenAI API Key Setup ---
openai_client = None
try:
    # Load the OpenAI API key from Streamlit secrets
    openai_api_key = st.secrets["OPENAI_API_KEY"]
    # Initialize the OpenAI client with the loaded API key
    openai_client = OpenAI(api_key=openai_api_key)
    st.success("OpenAI API key loaded successfully!")

except KeyError:
    st.error("OPENAI_API_KEY not found in Streamlit secrets! "
             "Please add your OpenAI API key to your app's secrets "
             "on Streamlit Community Cloud under the key 'OPENAI_API_KEY'.")
    st.stop()
except Exception as e:
    st.error(f"An unexpected error occurred during OpenAI API key setup: {e}")
    st.stop()

# --- Google Drive Reports Folder ID (from secrets) ---
GOOGLE_DRIVE_REPORTS_FOLDER_ID = None
try:
    GOOGLE_DRIVE_REPORTS_FOLDER_ID = st.secrets["GOOGLE_DRIVE_REPORTS_FOLDER_ID"]
except KeyError:
    st.error("GOOGLE_DRIVE_REPORTS_FOLDER_ID not found in Streamlit secrets! "
             "Please add the ID of your Google Drive reports folder to your app's secrets "
             "on Streamlit Community Cloud under the key 'GOOGLE_DRIVE_REPORTS_FOLDER_ID'.")
    st.stop()


# --- Utility Functions ---

# Function to hash passwords for Firestore storage (for local emulator login)
def hash_password(password):
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

# Function to check password against hash (for local emulator login)
def check_password(password, hashed_password):
    return bcrypt.checkpw(password.encode('utf-8'), hashed_password.encode('utf-8'))

def login_user(email, password): # Removed desired_login_type, as is_admin is from DB
    try:
        user = auth.get_user_by_email(email)
        user_doc_ref = db.collection('users').document(user.uid)
        user_doc = user_doc_ref.get()

        if user_doc.exists:
            user_data = user_doc.to_dict()
            hashed_password_from_db = user_data.get('hashed_password')
            is_admin_from_db = user_data.get('is_admin', False)
            username_from_db = user_data.get('username')
            has_set_username_from_db = user_data.get('has_set_username', False)

            if not hashed_password_from_db or not check_password(password, hashed_password_from_db):
                 st.error("Invalid credentials. Please check your password.")
                 return

            # Check if the attempted login type matches the user's actual role
            if st.session_state['is_admin_attempt'] and not is_admin_from_db:
                st.error("You attempted to log in as an Admin, but this account does not have admin privileges.")
                return
            if not st.session_state['is_admin_attempt'] and is_admin_from_db:
                st.warning("You logged in as a User, but this account has admin privileges. You can log in as Admin to access more features.")
                # Allow login but keep the current attempt as user. This might be desired.
                # If strict, would be: st.error("Please log in as Admin for this account.")
                # return

            st.session_state['logged_in'] = True
            st.session_state['user_email'] = email
            st.session_state['user_uid'] = user.uid
            st.session_state['is_admin'] = is_admin_from_db # Actual role from DB
            st.session_state['username'] = username_from_db
            st.session_state['has_set_username'] = has_set_username_from_db
            st.session_state['login_mode'] = 'logged_in' # Transition to logged-in state

            # Check if user needs to set username/password
            if not st.session_state['has_set_username']: 
                st.session_state['needs_username_setup'] = True
                st.success(f"Welcome, {email}! Please set up your display name and password.")
            else:
                st.session_state['needs_username_setup'] = False # Ensure this is false if already set
                if st.session_state['username']:
                    st.success(f"Logged in as {st.session_state['username']}!")
                else:
                    st.success(f"Logged in as {st.session_state['user_email']}!") # Fallback if username not set
            st.session_state['login_success'] = True
            
            if st.session_state['is_admin']:
                st.session_state['current_admin_page'] = 'generate' # Admin default to Generate Report
            else:
                st.session_state['current_admin_page'] = 'reports' # User default to Reports

            st.rerun() # Rerun to update UI after login
        else:
            st.error("User not found or password incorrect.")
    except auth.UserNotFoundError:
        st.error("User not found.")
    except Exception as e:
        st.error(f"Login error: {e}")

def create_user(email, password, is_admin=False):
    try:
        user_record = auth.create_user(email=email, password=password)
        user_ref = db.collection('users').document(user_record.uid)
        user_ref.set({
            'email': email,
            'is_admin': is_admin,
            'created_at': firestore.SERVER_TIMESTAMP,
            'hashed_password': hash_password(password), # Store hash for local emulator login
            'username': None, # New field: initially none
            'has_set_username': False # New flag: user needs to set username and password
        })
        st.success(f"User {email} created successfully!")
        return user_record.uid
    except exceptions.FirebaseError as e:
        error_message = e.code
        if "email-already-exists" in error_message:
            st.error("The email address is already in use by another account.")
        else:
            st.error(f"Error creating user: {error_message}")
        return None
    except Exception as e:
        st.error(f"An unexpected error occurred: {e}")
        return None

def logout_user():
    for key in ['logged_in', 'user_email', 'user_uid', 'is_admin', 'username', 'has_set_username', 'needs_username_setup', 'login_success', 'current_admin_page', 'login_mode', 'is_admin_attempt']:
        if key in st.session_state:
            del st.session_state[key]
    st.session_state['login_mode'] = 'choose_role' # Reset to role selection on logout
    st.rerun()

# --- Content Extraction Functions ---
def get_pdf_text(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

def get_docx_text(file):
    document = Document(file)
    text = ""
    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"
    return text

# --- OpenAI/AI Functions ---
def get_openai_response(prompt_text):
    # Use the globally initialized openai_client
    if openai_client:
        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o", # Using a powerful model
                messages=[
                    {"role": "system", "content": "You are a helpful AI assistant specialized in analyzing Job Descriptions and CVs. Provide concise, direct, and actionable insights. Be professional and objective."},
                    {"role": "user", "content": prompt_text}
                ],
                temperature=0.7 # Adjust creativity
            )
            return response.choices[0].message.content
        except Exception as e:
            st.error(f"Error calling OpenAI API: {e}. Please check your API key and network connection.")
            return "Error: Could not get response from AI."
    else:
        st.error("OpenAI client not initialized. Cannot generate AI response.")
        return "Error: OpenAI client not available."


# --- Report Generation Function ---
def create_comparative_docx_report(jd_text, cv_texts, report_data):
    document = Document()
    document.add_heading(report_data.get('report_title', 'JD-CV Comparative Analysis Report'), level=1)
    
    # Add a paragraph for general info
    document.add_paragraph(f"Generated by: {report_data.get('generated_by_username', report_data.get('generated_by_email'))}")
    document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"Job Description File: {report_data.get('jd_filename', 'N/A')}")
    document.add_paragraph(f"CV Files Analyzed: {', '.join(report_data.get('cv_filenames', ['N/A']))}")

    document.add_page_break()

    # JD Analysis Section
    document.add_heading('Job Description Analysis', level=2)
    jd_analysis_prompt = f"Analyze the following Job Description to identify key requirements, responsibilities, and qualifications. Structure the output clearly with bullet points:\n\nJD:\n{jd_text}"
    jd_analysis_response = get_openai_response(jd_analysis_prompt)
    document.add_paragraph(jd_analysis_response)
    document.add_paragraph('\n')

    # Overall CV Analysis (Concise Summary)
    document.add_heading('Overall CV Analysis (Summary)', level=2)
    overall_cv_prompt = f"Given the following Job Description and multiple CVs, provide an overall summary of how well the combined CVs generally align with the JD. Highlight common strengths and weaknesses across the candidates.\n\nJD:\n{jd_text}\n\nCVs:\n{'---CV---\n'.join(cv_texts)}"
    overall_cv_response = get_openai_response(overall_cv_prompt)
    document.add_paragraph(overall_cv_response)
    document.add_paragraph('\n')

    # Individual CV Comparison Section
    document.add_heading('Individual CV Comparison', level=2)
    for i, cv_text in enumerate(cv_texts):
        cv_filename = report_data['cv_filenames'][i] if i < len(report_data['cv_filenames']) else f"CV {i+1}"
        document.add_heading(f'{cv_filename} Comparison', level=3)
        prompt = f"Compare the following CV against the Job Description. Provide:\n1. Key strengths of the CV relative to the JD.\n2. Key areas of improvement/gaps in the CV relative to the JD.\n3. A concise overall fit score (e.g., 1-10 or Poor/Fair/Good/Excellent).\n\nJob Description:\n{jd_text}\n\nCandidate CV:\n{cv_text}"
        
        response = get_openai_response(prompt)
        document.add_paragraph(response)
        document.add_paragraph('\n')
        if i < len(cv_texts) - 1:
            document.add_page_break()

    # Save to buffer
    docx_buffer = io.BytesIO()
    document.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


# --- Streamlit Pages/Components ---

def choose_login_type_page():
    st.markdown("<div class='login-container'>", unsafe_allow_html=True)
    st.image("sso_logo.png", width=150)
    st.title("SSO Consultants")
    st.subheader("AI Recruitment Tool")
    st.markdown("---")
    st.markdown("### Choose Login Type")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("Login as User", key="user_login_button"): # Removed use_container_width
            st.session_state['login_mode'] = 'login_form'
            st.session_state['is_admin_attempt'] = False
            st.rerun()
    with col2:
        if st.button("Login as Admin", key="admin_login_button"): # Removed use_container_width
            st.session_state['login_mode'] = 'login_form'
            st.session_state['is_admin_attempt'] = True
            st.rerun()
    
    st.info("Note: New users can only be added by an administrator.")
    st.markdown("</div>", unsafe_allow_html=True)


def display_login_page():
    st.markdown("<div class='login-container'>", unsafe_allow_html=True)
    st.image("sso_logo.png", width=150)
    st.title("SSO Consultants")
    st.subheader("AI Recruitment Tool")
    st.markdown("---")

    login_title = "Admin Login" if st.session_state['is_admin_attempt'] else "User Login"
    st.markdown(f"### Please Log In as {login_title}")

    with st.form(key='login_form_main'):
        email = st.text_input("Email", key='email_input_login')
        password = st.text_input("Password", type="password", key='password_input_login')
        submit_button = st.form_submit_button(label='Log In')

        if submit_button:
            if email and password:
                login_user(email, password)
            else:
                st.warning("Please enter both email and password.")
    
    # Back button to choose role
    if st.button("Back to Login Choices", key="back_to_role_selection"):
        st.session_state['login_mode'] = 'choose_role'
        st.session_state['is_admin_attempt'] = False # Reset this flag
        st.rerun()

    st.info("Note: New users can only be added by an administrator.")
    st.markdown("</div>", unsafe_allow_html=True)


def generate_comparative_report_page():
    st.subheader("Generate Comparative Report")
    
    st.info("Upload a Job Description (JD) and one or more Candidate CVs (PDF or DOCX) to get a comparative analysis.", icon="ℹ️")

    # MODIFIED: Wrapped file uploaders in a custom left-aligned div
    st.markdown("<div class='left-aligned-content'>", unsafe_allow_html=True) 
    jd_file = st.file_uploader("Upload Job Description (PDF or DOCX)", type=["pdf", "docx"], key="jd_uploader")
    cv_files = st.file_uploader("Upload Candidate CVs (PDF or DOCX) - Multiple Allowed", type=["pdf", "docx"], accept_multiple_files=True, key="cv_uploader")
    st.markdown("</div>", unsafe_allow_html=True) # Close the custom div

    if st.button("Generate Comparative Report", key="generate_report_button"):
        if jd_file and cv_files:
            with st.spinner("Analyzing documents and generating report... This may take a moment."):
                jd_text = ""
                if jd_file.type == "application/pdf":
                    jd_text = get_pdf_text(jd_file)
                elif jd_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    jd_text = get_docx_text(jd_file)
                
                cv_texts = []
                for cv_file in cv_files:
                    if cv_file.type == "application/pdf":
                        cv_texts.append(get_pdf_text(cv_file))
                    elif cv_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        cv_texts.append(get_docx_text(cv_file))
                
                if not jd_text:
                    st.error("Could not extract text from Job Description. Please ensure it's a valid PDF/DOCX.")
                    return
                if not cv_texts:
                    st.error("Could not extract text from any CVs. Please ensure they are valid PDF/DOCX.")
                    return

                # Prepare report data for both generation and saving
                report_data = {
                    'report_title': f"Report for {jd_file.name}",
                    'generated_by_uid': st.session_state['user_uid'],
                    'generated_by_email': st.session_state['user_email'],
                    'generated_by_username': st.session_state['username'],
                    'jd_filename': jd_file.name,
                    'cv_filenames': [cv.name for cv in cv_files],
                }

                docx_buffer = create_comparative_docx_report(jd_text, cv_texts, report_data)
                
                # Generate unique filename for the report
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                report_full_filename = f"JD_CV_Analysis_Report_{timestamp}.docx"


                # --- Save report to Google Drive and Firestore ---
                if drive_service and GOOGLE_DRIVE_REPORTS_FOLDER_ID:
                    try:
                        file_metadata = {
                            'name': report_full_filename,
                            'parents': [GOOGLE_DRIVE_REPORTS_FOLDER_ID], # Specify the parent folder ID
                            'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        }
                        
                        docx_buffer.seek(0) # Ensure buffer is at the beginning for upload
                        media = MediaIoBaseUpload(docx_buffer,
                                                  mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                                  resumable=True)
                        
                        uploaded_file = drive_service.files().create(
                            body=file_metadata,
                            media_body=media,
                            fields='id,webViewLink' # Request ID and a viewable link
                        ).execute()

                        google_drive_file_id = uploaded_file.get('id')
                        google_drive_view_link = uploaded_file.get('webViewLink')

                        # Save report metadata to Firestore
                        db.collection('reports').add({
                            'report_title': report_data.get('report_title', 'JD-CV Comparative Analysis Report'),
                            'generated_by_uid': st.session_state['user_uid'],
                            'generated_by_email': st.session_state['user_email'],
                            'generated_by_username': st.session_state['username'],
                            'date_generated': firestore.SERVER_TIMESTAMP,
                            'jd_filename': jd_file.name,
                            'cv_filenames': [cv_file.name for cv_file in cv_files], # Use original file objects for names
                            'google_drive_file_id': google_drive_file_id, # Store the Drive file ID
                            'google_drive_view_link': google_drive_view_link, # Store the direct view link
                            'google_drive_folder_link': f"https://drive.google.com/drive/folders/{GOOGLE_DRIVE_REPORTS_FOLDER_ID}" # Link to the parent folder
                        })
                        st.success("Report saved to Google Drive and ready for download/view!")
                        docx_buffer.seek(0) # Reset buffer again for local download button

                    except Exception as e:
                        st.error(f"Error saving report to Google Drive: {e}")
                        st.warning("Report was generated but could not be saved to Google Drive. You can still download it locally.")
                else:
                    st.warning("Google Drive service not initialized or folder ID missing. Report not saved to cloud.")

                # Always provide local download option
                st.download_button(
                    label="Download Report Locally",
                    data=docx_buffer,
                    file_name=report_full_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_report_button"
                )
        else:
            st.warning("Please upload a Job Description and at least one CV to generate a report.")

def manage_users_page():
    st.subheader("Manage Users")

    # Display existing users
    st.markdown("### Existing Users")
    users_ref = db.collection('users').stream()
    users_data = []
    for user_doc in users_ref:
        user = user_doc.to_dict()
        users_data.append({
            "UID": user_doc.id,
            "Email": user.get('email', 'N/A'),
            "Username": user.get('username', 'Not Set'),
            "Admin": "Yes" if user.get('is_admin') else "No",
            "Created At": user.get('created_at').strftime('%Y-%m-%d %H:%M') if user.get('created_at') else 'N/A'
        })
    
    if users_data:
        users_df = pd.DataFrame(users_data)
        st.dataframe(users_df, use_container_width=True)
    else:
        st.info("No users found.")

    st.markdown("### Add New User")
    with st.form("add_user_form"):
        new_user_email = st.text_input("New User Email")
        new_user_password = st.text_input("Initial Password for New User", type="password")
        is_new_user_admin = st.checkbox("Grant Admin Privileges to New User?")
        
        proceed_with_creation = True
        if is_new_user_admin:
            st.warning("You are about to grant this user administrator privileges. This will give them full access to manage users, generate reports, and view all data. Please confirm this action.", icon="⚠️")
            confirm_admin_grant = st.checkbox("I understand and confirm to grant admin privileges", key="confirm_admin_grant")
            if not confirm_admin_grant:
                proceed_with_creation = False

        add_user_button = st.form_submit_button("Add User")

        if add_user_button:
            if new_user_email and new_user_password:
                if proceed_with_creation: # Only proceed if confirmed (or not admin user)
                    create_user(new_user_email, new_user_password, is_new_user_admin)
                    # The user list will update on the next interaction.
                else:
                    if is_new_user_admin: # Display specific warning if admin grant was not confirmed
                        st.error("Please confirm granting admin privileges to proceed.")
            else:
                st.warning("Please enter email and initial password for the new user.")


def show_all_reports_page():
    st.subheader("All Generated Reports")
    
    # Provide a direct link to the Google Drive reports folder
    if GOOGLE_DRIVE_REPORTS_FOLDER_ID:
        st.markdown(
            f"**All generated reports are stored in the following Google Drive folder:** "
            f"[View Reports Folder]({f'https://drive.google.com/drive/folders/{GOOGLE_DRIVE_REPORTS_FOLDER_ID}'})",
            unsafe_allow_html=True
        )
    else:
        st.warning("Google Drive Reports Folder ID is not configured in secrets.")
    st.write("---") # Separator

    reports_ref = db.collection('reports').order_by('date_generated', direction=firestore.Query.DESCENDING).stream()
    reports_data = []
    for report_doc in reports_ref:
        report = report_doc.to_dict()
        
        # Determine the link to use: Google Drive view link
        download_link = report.get('google_drive_view_link', '#') 
        
        reports_data.append({
            "Title": report.get('report_title', 'N/A'),
            "Generated By": report.get('generated_by_username', report.get('generated_by_email', 'N/A')),
            "Date": report.get('date_generated').strftime('%Y-%m-%d %H:%M') if report.get('date_generated') else 'N/A',
            "JD File": report.get('jd_filename', 'N/A'),
            "CV Files": ", ".join(report.get('cv_filenames', [])),
            "Link": f"[View Report]({download_link})" # Display as markdown link
        })
    
    if reports_data:
        reports_df = pd.DataFrame(reports_data)
        # Display as HTML to render markdown links
        st.markdown(reports_df.to_html(escape=False, index=False), unsafe_allow_html=True)
    else:
        st.info("No reports found.")


# --- Main Application Logic ---
def main():
    # Fixed top-left logo (outside the main content flow)
    try:
        with open("sso_logo.png", "rb") as f:
            logo_base64 = base64.b64encode(f.read()).decode()
        st.markdown(f'<img src="data:image/png;base64,{logo_base64}" class="fixed-top-left-logo">', unsafe_allow_html=True)
    except FileNotFoundError:
        st.warning("`sso_logo.png` not found for the fixed top-left logo. Please ensure it's in the root directory.")
    except Exception as e:
        st.error(f"Error loading fixed top-left logo: {e}")


    if not st.session_state['logged_in']:
        if st.session_state['login_mode'] == 'choose_role':
            choose_login_type_page()
        elif st.session_state['login_mode'] == 'login_form':
            display_login_page()
    else:
        # Handle username and password setup first if needed (for ALL new users)
        if st.session_state.get('needs_username_setup'): 
            st.header("Setup Your Account (First-Time Login)")
            st.warning("As this is your first login, please set a new display name and password for security.", icon="🔒")
            with st.form("set_account_details_form"):
                new_username = st.text_input("Enter your desired display name (e.g., Alex Smith)", key="new_username_input_setup")
                new_password = st.text_input("Set New Password", type="password", key="new_password_input_setup")
                confirm_password = st.text_input("Confirm New Password", type="password", key="confirm_password_input_setup")
                
                submit_setup_button = st.form_submit_button("Save Account Details and Re-Login")

                if submit_setup_button:
                    if not new_username:
                        st.warning("Please enter a display name.")
                    elif not new_password:
                        st.warning("Please set a new password.")
                    elif new_password != confirm_password:
                        st.error("Passwords do not match. Please re-enter.")
                    else:
                        try:
                            # 1. Update Firebase Authentication password
                            auth.update_user(st.session_state['user_uid'], password=new_password)
                            
                            # 2. Update Firestore for username and flags
                            user_doc_ref = db.collection('users').document(st.session_state['user_uid'])
                            user_doc_ref.update({
                                'username': new_username,
                                'has_set_username': True,
                                'hashed_password': hash_password(new_password) # IMPORTANT: Update hashed_password in Firestore for local emulator login validation
                            })
                            
                            st.session_state['username'] = new_username
                            st.session_state['has_set_username'] = True
                            st.session_state['needs_username_setup'] = False # This flag is now done

                            st.success(f"Account details updated successfully! Please log in with your new password.")
                            logout_user() # Log out to force re-login with new password
                            
                        except exceptions.FirebaseError as e:
                            st.error(f"Error updating account in Firebase: {e}")
                            st.info("Please ensure your Firebase setup is correct.")
                        except Exception as e:
                            st.error(f"An unexpected error occurred during account setup: {e}")
            return # Stop execution here until account details are set and user logs out
        
        # --- Sidebar Navigation for Logged-in Users ---
        with st.sidebar:
            st.image("sso_logo.png", use_container_width=True) 
            st.subheader(f"Welcome, {st.session_state['username'] if st.session_state['username'] else st.session_state['user_email']}!")
            st.write(f"Role: {'Admin' if st.session_state['is_admin'] else 'User'}")
            st.write("---")

            if st.session_state['is_admin']:
                st.write("### Admin Navigation")
                if st.button("Generate Report", key="nav_generate_admin"):
                    st.session_state['current_admin_page'] = 'generate'
                    st.rerun()
                if st.button("All Reports", key="nav_all_reports_admin"):
                    st.session_state['current_admin_page'] = 'reports'
                    st.rerun()
                if st.button("Manage Users", key="nav_manage_users_admin"):
                    st.session_state['current_admin_page'] = 'manage_users'
                    st.rerun()
            else:
                st.write("### User Navigation")
                if st.button("Generate Report", key="nav_generate_user"):
                    st.session_state['current_admin_page'] = 'generate'
                    st.rerun()
                if st.button("All Reports", key="nav_all_reports_user"):
                    st.session_state['current_admin_page'] = 'reports'
                    st.rerun()
            st.write("---")
            if st.button("Logout", key="logout_button"):
                logout_user()

        # --- Main Content Area ---
        if st.session_state['is_admin']:
            if st.session_state['current_admin_page'] == 'generate':
                generate_comparative_report_page()
            elif st.session_state['current_admin_page'] == 'reports':
                show_all_reports_page()
            elif st.session_state['current_admin_page'] == 'manage_users':
                manage_users_page()
        else: # Regular user view
            if st.session_state['current_admin_page'] == 'generate':
                generate_comparative_report_page()
            elif st.session_state['current_admin_page'] == 'reports':
                show_all_reports_page()


# --- Custom FOOTER (Always visible at the bottom of the page) ---
st.markdown(
    """
    <div style="
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        text-align: center;
        color: #4CAF50 !important; /* Green text for footer */
        padding: 10px;
        background-color: #FFFFFF; /* Match page background */
        font-size: 0.8em;
        border-top: 1px solid #E0E0E0; /* Subtle border for separation */
        z-index: 999;
    ">
        ©copyright SSO Consultants
    </div>
    """,
    unsafe_allow_html=True
)


if __name__ == "__main__":
    main()
